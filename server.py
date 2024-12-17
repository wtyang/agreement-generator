# -*- coding: utf-8 -*-
from flask import Flask, render_template, request, send_file
import os
from zipfile import ZipFile
from openpyxl import load_workbook
from docx import Document
from datetime import datetime


app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'upload'

def fill_contract(template_path, data, header):
    # header是从Excel文件的第一行读取的列名
    for row in data:
        document = Document(template_path)
        mapping = dict(zip(header, row))  # 创建一个字段名到数据的映射
        print(mapping)
        for paragraph in document.paragraphs:
            for key in mapping:
                if "temp_" + key + "" in paragraph.text:  # 查找占位符并替换
                    for run in paragraph.runs:
                        if "temp_" + key + "" in run.text:
                            run.text = run.text.replace("temp_" + key + "", str(mapping[key]))
                    # paragraph.text = paragraph.text.replace("temp_" + key + "", str(mapping[key]))
        for table in document.tables:
            print("table:",table)
            for row in table.rows:
                for cell in row.cells:
                    print("cell:",cell.text,"temp_" + key + "")
                    for key in mapping:
                        if "temp_" + key + "" in cell.text:  # 查找占位符并替换
                            print("          查找占位符并替换:",key,str(mapping[key]))
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    print("run:",run.text,"temp_" + key + "")
                                    if "temp_" + key + "" in run.text:
                                        print("替换:",run.text)
                                        run.text = run.text.replace("temp_" + key + "", str(mapping[key]))
        # 保存每份填充后的合同为一个新的Word文档
                    # 获取当前日期
        current_date = datetime.now()

        # 将日期格式化为"240322"
        formatted_date = current_date.strftime("%d%m%y")
        filled_contract_path = "{}-{}.docx".format(mapping[header[0]],formatted_date)  # 例如，以公司名称命名
        document.save(filled_contract_path)
        yield filled_contract_path  # 使用生成器返回路径，以便后续打包

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    """
    处理文件上传并生成打包文件。

    该端点期望接收两个文件：'data_file'（数据文件，如Excel）和'template_file'（模板文件）。
    函数将读取Excel数据，根据模板文件生成多个填充后的文件，并将它们打包成一个ZIP文件供下载。

    Returns:
        如果缺少任何文件，返回错误信息和400状态码。
        否则，返回一个ZIP文件，其中包含根据Excel数据填充的模板文件。
    """
    print(request.files)
    if 'data_file' not in request.files or 'template_file' not in request.files:
        return "缺少文件", 400
    data_file = request.files['data_file']
    template_file = request.files['template_file']

    data_filename = os.path.basename(data_file.filename)
    template_filename = os.path.basename(template_file.filename)
    print("数据文件名：{},模板文件名：{}".format(data_filename,template_filename))
    data_file_path = os.path.join(app.config['UPLOAD_FOLDER'], data_filename)
    template_file_path = os.path.join(app.config['UPLOAD_FOLDER'], template_filename)

    data_file.save(data_file_path)
    template_file.save(template_file_path)

    # 处理Excel数据
    wb = load_workbook(data_file_path)
    ws = wb.active
    data = []
    for row in ws.iter_rows(min_row=2, values_only=True):  # 从第二行开始读取数据，跳过表头
        data.append(row)
    header = [cell.value for cell in ws[1]]  # 读取表头
    # 打包数据
    zip_filename = "{}-打包.zip".format(template_filename)
    with ZipFile(zip_filename, 'w') as zipf:
        for filled_contract_path in fill_contract(template_file_path, data, header):
            zipf.write(filled_contract_path)  # 将每份填充后的合同添加到zip文件

    # 您可以将zip文件发送给用户以供下载。
    return send_file(zip_filename, as_attachment=True)

@app.route('/demoDownload')
def download_file():
    file_path = './template.zip'
    return send_file(file_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
